from __future__ import annotations

import re
from collections.abc import Mapping
from html import escape
from pathlib import Path


def register_chinese_pdf_font() -> str:
    """Register a ReportLab CID font available on every platform."""

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    font_name = "STSong-Light"
    if font_name not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    return font_name


def build_resume_pdf(
    resume: Mapping[str, object],
    output_path: str | Path,
) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    font_name = register_chinese_pdf_font()
    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=16 * mm,
    )
    title_style = ParagraphStyle(
        "ResumeTitle",
        fontName=font_name,
        fontSize=18,
        leading=24,
        textColor=colors.HexColor("#242638"),
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        fontName=font_name,
        fontSize=10.5,
        leading=17,
        textColor=colors.HexColor("#33384d"),
        spaceAfter=7,
    )
    title = escape(str(_value(resume, "title") or "简历"))
    content = str(_value(resume, "content") or "")
    story = [Paragraph(title, title_style)]
    for block in re.split(r"\n\s*\n", content):
        lines = "<br/>".join(escape(line) for line in block.splitlines())
        if lines.strip():
            story.append(Paragraph(lines, body_style))
            story.append(Spacer(1, 3))
    document.build(story)


def build_resume_docx(
    resume: Mapping[str, object],
    output_path: str | Path,
) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    document = Document()
    normal_font = document.styles["Normal"].font
    normal_font.name = "Arial"
    normal_font.size = Pt(10.5)
    normal_font._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    title = document.add_heading(
        str(_value(resume, "title") or "简历"),
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Arial"
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(36, 38, 56)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    document.add_paragraph()
    for block in re.split(
        r"\n\s*\n",
        str(_value(resume, "content") or ""),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.paragraph_format.space_after = Pt(6)
        for index, line in enumerate(block.splitlines()):
            if index:
                paragraph.add_run().add_break()
            run = paragraph.add_run(line)
            run.font.name = "Arial"
            run.font.size = Pt(10.5)
            run._element.rPr.rFonts.set(
                qn("w:eastAsia"),
                "Noto Sans CJK SC",
            )
    document.save(str(output_path))


def _value(resume: Mapping[str, object], key: str) -> object | None:
    getter = getattr(resume, "get", None)
    if callable(getter):
        return getter(key)
    try:
        return resume[key]
    except (IndexError, KeyError):
        return None
