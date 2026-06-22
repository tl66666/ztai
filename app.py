from __future__ import annotations

from datetime import datetime
from html import escape
import json
import os
import re
import shutil
import sqlite3
import subprocess
import tempfile
import uuid

from flask import Flask, jsonify, request, send_file, send_from_directory
from flask_cors import CORS

from utils.ai_client import extract_keywords, get_ai_client, set_api_key


BASE_DIR = os.path.dirname(__file__)
DB_PATH = os.environ.get("JOBHUNTER_DB_PATH", os.path.join(BASE_DIR, "jobhunter.db"))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
EXPORT_FOLDER = os.path.join(BASE_DIR, "exports")
ALLOWED_EXTENSIONS = {"pdf", "doc", "docx", "txt", "png", "jpg", "jpeg"}

app = Flask(__name__, static_folder="static")
CORS(app)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(EXPORT_FOLDER, exist_ok=True)


INTERVIEW_SESSIONS: dict[str, dict] = {}


class ClosingConnection(sqlite3.Connection):
    def __exit__(self, exc_type, exc_value, traceback):
        result = super().__exit__(exc_type, exc_value, traceback)
        self.close()
        return result


def get_db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH, factory=ClosingConnection)
    conn.row_factory = sqlite3.Row
    return conn


def init_db() -> None:
    with get_db() as conn:
        conn.executescript(
            """
            CREATE TABLE IF NOT EXISTS resumes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                title TEXT NOT NULL,
                content TEXT NOT NULL,
                file_path TEXT,
                file_type TEXT,
                analysis_result TEXT,
                tailored_result TEXT,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT DEFAULT CURRENT_TIMESTAMP
            );

            CREATE TABLE IF NOT EXISTS job_matches (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                resume_id INTEGER NOT NULL,
                job_title TEXT NOT NULL,
                match_score INTEGER,
                analysis TEXT,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            );

            CREATE TABLE IF NOT EXISTS interviews (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                resume_id INTEGER,
                job_title TEXT NOT NULL,
                conversation TEXT,
                score INTEGER,
                feedback TEXT,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            );

            CREATE TABLE IF NOT EXISTS job_applications (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                company TEXT NOT NULL,
                job_title TEXT NOT NULL,
                status TEXT DEFAULT '已投递',
                city TEXT,
                salary_min INTEGER,
                salary_max INTEGER,
                notes TEXT,
                applied_at TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT DEFAULT CURRENT_TIMESTAMP
            );

            CREATE TABLE IF NOT EXISTS practice_records (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER DEFAULT 1,
                category TEXT,
                question TEXT NOT NULL,
                answer TEXT,
                score INTEGER,
                feedback TEXT,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            );

            CREATE TABLE IF NOT EXISTS audio_records (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER DEFAULT 1,
                transcript TEXT NOT NULL,
                audio_file TEXT,
                score INTEGER,
                metrics TEXT,
                feedback TEXT,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            );
            """
        )
        ensure_column(conn, "resumes", "analysis_result", "TEXT")
        ensure_column(conn, "resumes", "tailored_result", "TEXT")


def ensure_column(conn: sqlite3.Connection, table: str, column: str, column_type: str) -> None:
    columns = {row["name"] for row in conn.execute(f"PRAGMA table_info({table})").fetchall()}
    if column not in columns:
        conn.execute(f"ALTER TABLE {table} ADD COLUMN {column} {column_type}")


def row_to_dict(row: sqlite3.Row | None) -> dict | None:
    return dict(row) if row else None


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def parse_resume_file(file_path: str, file_type: str) -> str:
    try:
        if file_type == "txt":
            for encoding in ("utf-8", "gbk"):
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
        if file_type == "pdf":
            import PyPDF2

            text = []
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text.append(page.extract_text() or "")
            return "\n".join(text).strip()
        if file_type in {"doc", "docx"}:
            from docx import Document

            doc = Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return "图片简历已上传。建议手动补充文本内容，便于 AI 分析。"
    except Exception as exc:
        return f"文件已上传，但解析失败：{exc}"


def get_resume_or_404(resume_id: int):
    conn = get_db()
    row = conn.execute("SELECT * FROM resumes WHERE id = ?", (resume_id,)).fetchone()
    conn.close()
    return row


def safe_filename(name: str, suffix: str = "") -> str:
    stem = re.sub(r"[^\w\u4e00-\u9fa5.-]+", "_", name or "resume").strip("._") or "resume"
    return f"{stem}{suffix}"


def audio_mime(filename: str) -> str:
    ext = os.path.splitext(filename.lower())[1]
    return {
        ".mp3": "audio/mpeg",
        ".wav": "audio/wav",
        ".webm": "audio/webm",
        ".ogg": "audio/ogg",
        ".m4a": "audio/mp4",
    }.get(ext, "application/octet-stream")


def convert_audio_file(source_path: str, target_format: str) -> tuple[str, str]:
    target_format = target_format.lower()
    if target_format not in {"mp3", "wav"}:
        raise ValueError("Unsupported audio format")
    ffmpeg = shutil.which("ffmpeg")
    if not ffmpeg:
        raise RuntimeError("当前环境未安装 ffmpeg，无法转码为 MP3/WAV")
    stem = os.path.splitext(os.path.basename(source_path))[0]
    target_name = safe_filename(stem, f".{target_format}")
    target_path = os.path.join(EXPORT_FOLDER, target_name)
    command = [ffmpeg, "-y", "-i", source_path, "-vn"]
    if target_format == "mp3":
        command += ["-codec:a", "libmp3lame", "-b:a", "192k"]
    else:
        command += ["-acodec", "pcm_s16le", "-ar", "44100", "-ac", "2"]
    command.append(target_path)
    subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    return target_path, target_name


def register_chinese_pdf_font() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
    ]
    for font_path in font_candidates:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont("JobHunterCN", font_path))
                return "JobHunterCN"
            except Exception:
                continue
    return "Helvetica"


def build_resume_pdf(resume: sqlite3.Row, output_path: str) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    font_name = register_chinese_pdf_font()
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=16 * mm,
    )
    title_style = ParagraphStyle(
        "ResumeTitle",
        fontName=font_name,
        fontSize=18,
        leading=24,
        textColor=colors.HexColor("#242638"),
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        fontName=font_name,
        fontSize=10.5,
        leading=17,
        textColor=colors.HexColor("#33384d"),
        spaceAfter=7,
    )
    story = [Paragraph(escape(resume["title"]), title_style)]
    for block in re.split(r"\n\s*\n", resume["content"] or ""):
        lines = "<br/>".join(escape(line) for line in block.splitlines())
        if lines.strip():
            story.append(Paragraph(lines, body_style))
            story.append(Spacer(1, 3))
    doc.build(story)


def build_resume_docx(resume: sqlite3.Row, output_path: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    title = document.add_heading(resume["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(36, 38, 56)
    document.add_paragraph()
    for block in re.split(r"\n\s*\n", resume["content"] or ""):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.paragraph_format.space_after = Pt(6)
        for index, line in enumerate(block.splitlines()):
            if index:
                paragraph.add_run().add_break()
            run = paragraph.add_run(line)
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(10.5)
    document.save(output_path)


def convert_word_to_pdf_native(source_path: str, target_path: str) -> bool:
    abs_source = os.path.abspath(source_path)
    abs_target = os.path.abspath(target_path)
    try:
        import pythoncom
        import win32com.client

        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(abs_source, ReadOnly=True)
        doc.SaveAs(abs_target, FileFormat=17)
        doc.Close(False)
        word.Quit()
        pythoncom.CoUninitialize()
        return os.path.exists(abs_target) and os.path.getsize(abs_target) > 0
    except Exception:
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        out_dir = os.path.dirname(abs_target)
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, abs_source],
            capture_output=True,
            text=True,
            timeout=60,
        )
        generated = os.path.join(out_dir, f"{os.path.splitext(os.path.basename(abs_source))[0]}.pdf")
        if result.returncode == 0 and os.path.exists(generated):
            if generated != abs_target:
                shutil.move(generated, abs_target)
            return True
    return False


def convert_pdf_to_word_file(source_path: str, target_path: str) -> None:
    from pdf2docx import Converter

    converter = Converter(source_path)
    try:
        converter.convert(target_path, start=0, end=None)
    finally:
        converter.close()


def score_resume_against_jd(resume_text: str, jd: str) -> tuple[int, list[str], list[str]]:
    resume_keywords = set(extract_keywords(resume_text))
    jd_keywords = set(extract_keywords(jd))
    if not jd_keywords:
        jd_keywords = {"项目", "沟通", "学习", "测试", "开发"}
    matched = sorted(resume_keywords & jd_keywords)
    missing = sorted(jd_keywords - resume_keywords)
    score = int(min(96, max(35, 50 + len(matched) * 8 - len(missing) * 3)))
    return score, matched, missing


def tailor_resume_locally(resume_text: str, job_title: str, jd: str) -> dict:
    score, matched, missing = score_resume_against_jd(resume_text, jd)
    top_keywords = matched + missing[:4]
    audit = build_resume_audit(resume_text, job_title, jd)
    positioning = f"面向 {job_title or '目标岗位'} 的候选人定位：具备项目实践、测试/开发工具链和 AI 应用理解，适合强调落地能力。"
    jd_focus = extract_jd_focus(jd)
    rewritten = [
        f"求职意向：{job_title or '目标岗位'}",
        "个人优势：具备 AI Web 系统项目实践，能围绕真实求职流程完成需求拆解、功能验证、接口验证和体验优化。",
        "",
        "项目经历改写示例：AI 智能体求职辅助 Web 系统",
        "- 项目背景：面向应届生/转岗求职者，设计简历优化、JD 匹配、模拟面试、语音表达分析和投递追踪的一站式求职辅助系统。",
        "- 个人职责：负责核心功能测试与体验优化，围绕简历上传解析、模型 API 兜底、JD 匹配分、面试状态流转设计测试场景。",
        "- 技术与工具：Flask、SQLite、JavaScript、Selenium、Postman、JMeter、Pytest，覆盖功能测试、接口测试和基础性能验证。",
        "- 结果产出：沉淀测试用例、缺陷记录、测试总结和可演示系统，将课程实训成果包装为可写入简历的完整项目经历。",
        "",
        "可量化表达模板：",
        "- 将“参与测试”改为“设计 X 类核心用例，覆盖上传、匹配、面试、看板等主流程，发现并推动修复 X 个问题”。",
        "- 将“熟悉工具”改为“使用 JMeter 对关键接口进行并发压测，记录响应时间、错误率和瓶颈结论”。",
        "",
        "原始简历摘要：",
        resume_text[:900],
    ]
    return {
        "positioning": positioning,
        "match_score": score,
        "score_detail": audit["section_scores"],
        "brutal_comments": audit["brutal_comments"],
        "evidence_gaps": audit["evidence_gaps"],
        "jd_focus": jd_focus,
        "matched_keywords": matched,
        "keyword_gaps": missing,
        "keyword_strategy": top_keywords,
        "tailored_resume": "\n".join(rewritten),
        "rewrite_tips": [
            "每段项目经历补充使用工具、负责动作、验证对象和结果。",
            "把“参与项目”改成“负责模块/设计用例/定位问题/输出报告”。",
            "关键词不要堆砌，放进真实项目语境里。",
        ],
        "interview_talking_points": [
            "为什么做这个系统：从真实求职痛点出发，解决简历、岗位、面试准备割裂的问题。",
            "技术亮点：多模型路由、本地兜底、面试状态机、语音表达指标分析。",
            "测试亮点：围绕主流程、异常输入、接口返回、模型不可用场景设计用例。",
        ],
    }


def extract_jd_focus(jd: str) -> dict:
    text = jd or ""
    focus_map = {
        "硬技能": ["Python", "Java", "Flask", "Vue", "React", "MySQL", "Redis", "Docker", "Selenium", "JMeter", "Postman", "Pytest"],
        "测试能力": ["功能测试", "接口测试", "自动化测试", "性能测试", "测试用例", "缺陷", "回归测试"],
        "AI 能力": ["AI", "智能体", "大模型", "Prompt", "模型", "算法"],
        "软技能": ["沟通", "协作", "推动", "学习", "文档", "总结"],
    }
    return {name: [word for word in words if word.lower() in text.lower()] for name, words in focus_map.items()}


def analyze_voice_text(answer: str, duration_seconds: float | None = None, audio_metrics: dict | None = None) -> dict:
    clean = answer.strip()
    cn_chars = len(re.findall(r"[\u4e00-\u9fa5]", clean))
    words = re.findall(r"[A-Za-z]+", clean)
    unit_count = cn_chars + len(words)
    duration = max(float(duration_seconds or 0), unit_count / 3.2, 20)
    speed = round(unit_count / duration * 60)
    fillers = ["嗯", "呃", "啊", "然后", "就是", "那个", "这个", "的话", "其实", "可能"]
    filler_detail = {word: clean.count(word) for word in fillers if clean.count(word)}
    filler_count = sum(filler_detail.values())
    structure_markers = ["首先", "其次", "最后", "背景", "任务", "行动", "结果", "因为", "所以"]
    structure_score = sum(1 for marker in structure_markers if marker in clean)
    tips = []
    if speed > 260:
        tips.append("语速偏快，建议关键经历处主动停顿 0.5 秒，让面试官跟上信息。")
    elif speed < 140:
        tips.append("语速偏慢，可以提前准备 2 分钟版本，减少犹豫停顿。")
    else:
        tips.append("语速处在较自然区间，继续保持。")
    if filler_count > 2:
        tips.append(f"口头禅出现 {filler_count} 次，优先减少“然后、就是、那个”。")
    if structure_score < 2:
        tips.append("结构感不足，建议使用 STAR 或“背景-行动-结果”组织答案。")
    if unit_count < 60:
        tips.append("回答略短，需要补充具体项目细节和量化结果。")
    metrics = audio_metrics or {}
    audio_quality = "未提供真实音频，仅按文本估算表达表现。"
    if metrics:
        silence_ratio = float(metrics.get("silence_ratio") or 0)
        clipping_ratio = float(metrics.get("clipping_ratio") or 0)
        average_volume = float(metrics.get("average_volume") or 0)
        audio_parts = []
        if silence_ratio > 0.55:
            audio_parts.append("停顿偏多，像在临时组织语言")
            tips.append("录音停顿比例偏高，建议提前准备“结论-证据-结果”三句骨架。")
        elif silence_ratio < 0.15:
            audio_parts.append("停顿较少，表达连贯但要注意给面试官消化时间")
        else:
            audio_parts.append("停顿比例较自然")
        if clipping_ratio > 0.02:
            audio_parts.append("存在爆音/过载")
            tips.append("录音出现爆音，正式面试前调整麦克风距离和系统音量。")
        if average_volume and average_volume < 0.025:
            audio_parts.append("音量偏小，面试时可能听不清")
            tips.append("音量偏小，建议提高麦克风增益或靠近一点。")
        elif average_volume > 0.22:
            audio_parts.append("音量偏大，注意不要贴麦")
        audio_quality = "；".join(audio_parts)
    pace_label = "偏快" if speed > 260 else ("偏慢" if speed < 140 else "自然")
    filler_ratio = round(filler_count / max(1, unit_count) * 100, 2)
    keyword_hits = [word for word in ["项目", "测试", "接口", "自动化", "性能", "结果", "推动", "优化", "用户"] if word in clean]
    return {
        "overall_score": max(35, min(95, 82 - filler_count * 4 + structure_score * 3 - abs(speed - 200) // 12)),
        "estimated_speech_rate": speed,
        "pace_label": pace_label,
        "filler_count": filler_count,
        "filler_ratio": filler_ratio,
        "filler_detail": filler_detail,
        "structure_score": structure_score,
        "keyword_hits": keyword_hits,
        "clarity": "清晰" if structure_score >= 2 else "需要加强",
        "audio_quality": audio_quality,
        "audio_metrics": metrics,
        "dimension_scores": {
            "表达流畅": max(40, min(95, 90 - filler_count * 6)),
            "结构逻辑": max(35, min(95, 55 + structure_score * 10)),
            "岗位相关": max(35, min(95, 50 + len(keyword_hits) * 7)),
            "信息密度": max(35, min(95, min(90, unit_count))),
        },
        "tips": tips,
    }


def question_bank() -> dict:
    return {
        "general": [
            {"question": "请做一个 2 分钟自我介绍。", "answer": "用 当前身份 + 目标岗位 + 1 个核心项目 + 2 个能力证据 + 求职动机 收束。"},
            {"question": "你为什么选择这个岗位？", "answer": "从兴趣、能力匹配、项目经历、长期成长四点回答。"},
            {"question": "你有什么问题想问我？", "answer": "问岗位挑战、团队技术栈、入职前三个月期待，避免一上来只问福利。"},
            {"question": "你最大的优势是什么？", "answer": "选择与岗位相关的优势，用项目证据支撑，例如测试细致、学习快、能推动问题闭环。"},
            {"question": "你最大的缺点是什么？", "answer": "选择非致命缺点，说明改进动作和结果，不要说与岗位核心能力冲突的缺点。"},
        ],
        "test": [
            {"question": "如何设计 Web 系统的测试用例？", "answer": "按业务流程、输入边界、异常场景、权限、兼容性、性能和接口契约拆分。"},
            {"question": "接口测试重点关注什么？", "answer": "关注状态码、响应结构、业务字段、幂等性、鉴权、异常参数和数据一致性。"},
            {"question": "JMeter 性能测试怎么看结果？", "answer": "看吞吐量、平均/中位/P95 响应时间、错误率、资源瓶颈和并发拐点。"},
            {"question": "发现一个偶现 bug 你怎么处理？", "answer": "先记录环境和复现路径，再补日志、缩小变量、固定数据、提高复现概率，最后给出证据链。"},
            {"question": "自动化测试适合覆盖哪些场景？", "answer": "适合稳定、高频、回归成本高的主流程，不适合频繁变化和强视觉主观判断场景。"},
        ],
        "frontend": [
            {"question": "如何提升前端页面可用性？", "answer": "从信息架构、操作反馈、加载状态、表单校验、响应式和无障碍入手。"},
            {"question": "前端如何处理接口异常？", "answer": "区分网络错误、业务错误、超时和空数据，给出可恢复操作和明确提示。"},
        ],
        "python": [
            {"question": "Flask 项目如何组织接口？", "answer": "按业务模块拆分路由、服务、数据访问和配置，统一错误处理与响应结构。"},
            {"question": "SQLite 在小型项目里适合什么场景？", "answer": "适合课程项目、单机演示和轻量数据存储，部署简单，但高并发和复杂权限场景应换 MySQL/PostgreSQL。"},
        ],
        "ai": [
            {"question": "多模型接入为什么要做本地兜底？", "answer": "因为 API Key、网络、限流都可能失败，本地兜底能保证核心流程可演示、可测试、可用。"},
            {"question": "AI Agent 和普通聊天接口有什么区别？", "answer": "Agent 需要有目标、上下文、工具调用和流程状态，不只是单轮问答。"},
        ],
    }


@app.route("/")
def index():
    return send_from_directory("static", "index.html")


@app.route("/<path:path>")
def static_files(path):
    return send_from_directory("static", path)


@app.route("/api/uploads/<path:filename>")
def uploaded_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)


@app.route("/api/uploads/<path:filename>/download/<format_type>")
def download_audio_file(filename, format_type):
    source_path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(source_path):
        return jsonify({"success": False, "message": "音频文件不存在或已被删除"}), 404
    format_type = (format_type or "original").lower()
    if format_type == "original":
        return send_file(
            source_path,
            as_attachment=True,
            download_name=os.path.basename(filename),
            mimetype=audio_mime(filename),
        )
    try:
        target_path, target_name = convert_audio_file(source_path, format_type)
        return send_file(target_path, as_attachment=True, download_name=target_name, mimetype=audio_mime(target_name))
    except RuntimeError as exc:
        return jsonify({"success": False, "message": str(exc)}), 501
    except Exception as exc:
        return jsonify({"success": False, "message": f"音频转码失败：{exc}"}), 500


@app.route("/api/config/providers")
def providers():
    client = get_ai_client()
    return jsonify({"success": True, "providers": client.available_providers(), "active_provider": client.provider.id})


@app.route("/api/config/ai-key", methods=["POST"])
def configure_ai_key():
    data = request.get_json() or {}
    provider_id = data.get("provider", "glm")
    model_id = data.get("model") or data.get("model_id") or ""
    api_key = data.get("api_key", "")
    client = set_api_key(api_key, provider_id, model_id)
    return jsonify({"success": True, "provider": client.provider.id, "model": client.model, "ai_enabled": bool(api_key)})


@app.route("/api/config/ai-status")
def ai_status():
    client = get_ai_client()
    return jsonify({
        "success": True,
        "ai_enabled": bool(client.api_key),
        "provider": client.provider.id,
        "provider_name": client.provider.name,
        "model": client.model,
        "selected_model": client.model,
        "providers": client.available_providers(),
    })


@app.route("/api/resumes", methods=["POST"])
def create_resume():
    if request.files:
        file = request.files.get("file")
        user_id = request.form.get("user_id", 1)
        title = (request.form.get("title") or (file.filename if file else "未命名简历")).strip()
        if not file or not file.filename or not allowed_file(file.filename):
            return jsonify({"success": False, "message": "请上传 PDF、Word、TXT 或图片格式简历。"}), 400
        file_type = file.filename.rsplit(".", 1)[1].lower()
        filename = f"resume_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{file.filename}"
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        content = parse_resume_file(file_path, file_type)
    else:
        data = request.get_json() or {}
        user_id = data.get("user_id", 1)
        title = (data.get("title") or "").strip()
        content = (data.get("content") or "").strip()
        file_path = None
        file_type = None
        if not title or not content:
            return jsonify({"success": False, "message": "标题和内容不能为空。"}), 400

    with get_db() as conn:
        cursor = conn.execute(
            "INSERT INTO resumes (user_id, title, content, file_path, file_type) VALUES (?, ?, ?, ?, ?)",
            (user_id, title, content, file_path, file_type),
        )
        resume_id = cursor.lastrowid
    return jsonify({"success": True, "message": "简历已保存", "resume_id": resume_id, "parsed_content": content[:1000]}), 201


@app.route("/api/resumes/upload", methods=["POST"])
def upload_resume():
    return create_resume()


@app.route("/api/resumes/<int:user_id>")
def list_resumes(user_id):
    with get_db() as conn:
        rows = conn.execute("SELECT * FROM resumes WHERE user_id = ? ORDER BY updated_at DESC", (user_id,)).fetchall()
    return jsonify({"success": True, "data": [dict(row) for row in rows]})


@app.route("/api/resumes/detail/<int:resume_id>")
def resume_detail(resume_id):
    row = get_resume_or_404(resume_id)
    if not row:
        return jsonify({"success": False, "message": "简历不存在"}), 404
    return jsonify({"success": True, "data": dict(row)})


@app.route("/api/resumes/<int:resume_id>/original")
def resume_original(resume_id):
    row = get_resume_or_404(resume_id)
    if not row:
        return jsonify({"success": False, "message": "简历不存在"}), 404
    original_path = row["file_path"]
    if not original_path or not os.path.exists(original_path):
        return jsonify({"success": False, "message": "这份简历没有保存原始文件，只能编辑文本内容。"}), 404
    return send_file(original_path, as_attachment=False, download_name=os.path.basename(original_path))


@app.route("/api/resumes/<int:resume_id>/replace-file", methods=["POST"])
def replace_resume_file(resume_id):
    row = get_resume_or_404(resume_id)
    if not row:
        return jsonify({"success": False, "message": "简历不存在"}), 404
    file = request.files.get("file")
    if not file or not file.filename or not allowed_file(file.filename):
        return jsonify({"success": False, "message": "请上传 PDF、Word、TXT 或图片格式简历。"}), 400
    file_type = file.filename.rsplit(".", 1)[1].lower()
    filename = f"resume_replace_{resume_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{file.filename}"
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(file_path)
    content = parse_resume_file(file_path, file_type)
    with get_db() as conn:
        conn.execute(
            "UPDATE resumes SET file_path = ?, file_type = ?, content = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            (file_path, file_type, content, resume_id),
        )
    return jsonify({"success": True, "message": "原文件已替换并重新解析", "parsed_content": content[:1000]})


@app.route("/api/resumes/<int:resume_id>", methods=["PUT"])
def update_resume(resume_id):
    data = request.get_json() or {}
    title = (data.get("title") or "").strip()
    content = (data.get("content") or "").strip()
    if not title or not content:
        return jsonify({"success": False, "message": "标题和内容不能为空"}), 400
    with get_db() as conn:
        cursor = conn.execute(
            "UPDATE resumes SET title = ?, content = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            (title, content, resume_id),
        )
    if cursor.rowcount == 0:
        return jsonify({"success": False, "message": "简历不存在"}), 404
    return jsonify({"success": True, "message": "简历已更新"})


@app.route("/api/resumes/<int:resume_id>", methods=["DELETE"])
def delete_resume(resume_id):
    with get_db() as conn:
        cursor = conn.execute("DELETE FROM resumes WHERE id = ?", (resume_id,))
    if cursor.rowcount == 0:
        return jsonify({"success": False, "message": "简历不存在"}), 404
    return jsonify({"success": True, "message": "简历已删除"})


@app.route("/api/resumes/<int:resume_id>/export/<format_type>")
def export_resume(resume_id, format_type):
    row = get_resume_or_404(resume_id)
    if not row:
        return jsonify({"success": False, "message": "简历不存在"}), 404
    fmt = format_type.lower()
    if fmt not in {"pdf", "word", "docx"}:
        return jsonify({"success": False, "message": "仅支持 PDF 或 Word 导出"}), 400

    extension = "pdf" if fmt == "pdf" else "docx"
    filename = safe_filename(row["title"], f".{extension}")
    output_path = os.path.join(EXPORT_FOLDER, f"{uuid.uuid4().hex}_{filename}")

    original_path = row["file_path"]
    original_type = (row["file_type"] or "").lower()
    if original_path and os.path.exists(original_path):
        if extension == "docx" and original_type in {"doc", "docx"}:
            return send_file(original_path, as_attachment=True, download_name=filename)
        if extension == "pdf" and original_type == "pdf":
            return send_file(original_path, as_attachment=True, download_name=filename)
        if extension == "pdf" and original_type in {"doc", "docx"}:
            if convert_word_to_pdf_native(original_path, output_path):
                return send_file(output_path, as_attachment=True, download_name=filename)
        if extension == "docx" and original_type == "pdf":
            try:
                convert_pdf_to_word_file(original_path, output_path)
                return send_file(output_path, as_attachment=True, download_name=filename)
            except Exception:
                pass

    if extension == "pdf":
        build_resume_pdf(row, output_path)
    else:
        build_resume_docx(row, output_path)
    return send_file(output_path, as_attachment=True, download_name=filename)


@app.route("/api/convert/pdf-to-word", methods=["POST"])
def convert_pdf_to_word():
    file = request.files.get("file")
    if not file or not file.filename.lower().endswith(".pdf"):
        return jsonify({"success": False, "message": "请上传 PDF 文件"}), 400
    with tempfile.TemporaryDirectory() as temp_dir:
        source = os.path.join(temp_dir, safe_filename(file.filename))
        target_name = safe_filename(os.path.splitext(file.filename)[0], ".docx")
        target = os.path.join(EXPORT_FOLDER, f"{uuid.uuid4().hex}_{target_name}")
        file.save(source)
        try:
            convert_pdf_to_word_file(source, target)
        except Exception as exc:
            return jsonify({"success": False, "message": f"PDF 转 Word 失败：{exc}"}), 500
    return send_file(target, as_attachment=True, download_name=target_name)


@app.route("/api/convert/word-to-pdf", methods=["POST"])
def convert_word_to_pdf():
    file = request.files.get("file")
    if not file or not file.filename.lower().endswith((".doc", ".docx")):
        return jsonify({"success": False, "message": "请上传 Word 文件"}), 400
    with tempfile.TemporaryDirectory() as temp_dir:
        source = os.path.join(temp_dir, safe_filename(file.filename))
        target_name = safe_filename(os.path.splitext(file.filename)[0], ".pdf")
        target = os.path.join(EXPORT_FOLDER, f"{uuid.uuid4().hex}_{target_name}")
        file.save(source)
        try:
            if not convert_word_to_pdf_native(source, target):
                from docx import Document

                document = Document(source)
                text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
                build_resume_pdf({"title": os.path.splitext(file.filename)[0], "content": text}, target)
        except Exception as exc:
            return jsonify({"success": False, "message": f"Word 转 PDF 失败：{exc}"}), 500
    return send_file(target, as_attachment=True, download_name=target_name)


@app.route("/api/resumes/<int:resume_id>/analyze", methods=["POST"])
def analyze_resume(resume_id):
    row = get_resume_or_404(resume_id)
    if not row:
        return jsonify({"success": False, "message": "简历不存在"}), 404
    data = request.get_json() or {}
    result = get_ai_client().analyze_resume(row["content"], data.get("job_title", ""))
    keywords = extract_keywords(row["content"])
    analysis = result["content"]
    with get_db() as conn:
        conn.execute("UPDATE resumes SET analysis_result = ? WHERE id = ?", (analysis, resume_id))
    return jsonify({"success": True, "analysis": analysis, "keywords": keywords, "ai_used": result["success"], "provider": result["provider"]})


@app.route("/api/resumes/<int:resume_id>/audit", methods=["POST"])
def audit_resume(resume_id):
    row = get_resume_or_404(resume_id)
    if not row:
        return jsonify({"success": False, "message": "简历不存在"}), 404
    data = request.get_json() or {}
    audit = build_resume_audit(row["content"], data.get("job_title", ""), data.get("jd", ""))
    return jsonify({"success": True, **audit})


@app.route("/api/resumes/<int:resume_id>/improve", methods=["POST"])
def improve_resume(resume_id):
    row = get_resume_or_404(resume_id)
    if not row:
        return jsonify({"success": False, "message": "简历不存在"}), 404
    data = request.get_json() or {}
    job_title = data.get("job_title") or "目标岗位"
    jd = data.get("jd") or ""
    improved = build_improved_resume(row["content"], job_title, jd)
    new_title = f"{row['title']}-优化版"
    new_id = None
    if data.get("save", True):
        with get_db() as conn:
            cursor = conn.execute(
                "INSERT INTO resumes (user_id, title, content, analysis_result, tailored_result) VALUES (?, ?, ?, ?, ?)",
                (
                    row["user_id"],
                    new_title,
                    improved["improved_resume"],
                    json.dumps(improved["audit"], ensure_ascii=False),
                    json.dumps(improved, ensure_ascii=False),
                ),
            )
            new_id = cursor.lastrowid
    return jsonify({"success": True, "new_resume_id": new_id, "new_title": new_title, **improved})


def build_resume_audit(resume_text: str, job_title: str = "", jd: str = "") -> dict:
    text = resume_text or ""
    lower_text = text.lower()
    keywords = extract_keywords(text)
    score, matched, missing = score_resume_against_jd(text, jd or job_title)
    project_words = ["项目", "系统", "平台", "模块", "接口", "测试", "数据库", "前端", "后端", "模型", "智能体"]
    tool_words = ["flask", "sqlite", "python", "javascript", "vue", "react", "mysql", "redis", "docker", "git", "postman", "jmeter", "selenium", "pytest"]
    action_words = ["负责", "设计", "实现", "优化", "分析", "定位", "验证", "推动", "输出", "沉淀", "完成"]
    result_words = ["提升", "降低", "覆盖", "发现", "修复", "通过", "稳定", "响应", "效率", "错误率", "缺陷", "结果"]
    structure_words = ["背景", "目标", "职责", "行动", "结果", "产出", "复盘", "问题", "方案"]
    has_project = any(word in text for word in project_words)
    has_metrics = bool(re.search(r"\d+|%|次|条|个|ms|秒|分钟|小时|天|qps|接口|用例", text, re.I))
    matched_tools = [word for word in tool_words if word in lower_text]
    action_count = sum(text.count(word) for word in action_words)
    result_count = sum(text.count(word) for word in result_words)
    structure_count = sum(text.count(word) for word in structure_words)
    length_score = min(18, len(text) // 80)
    project_score = 10 + min(16, sum(text.count(word) for word in project_words) * 2)
    metric_score = 8 + (14 if has_metrics else 0)
    action_score = 8 + min(14, action_count * 3)
    tool_score = 8 + min(14, len(matched_tools) * 3)
    jd_score = max(6, min(18, score // 5)) if jd else 10
    structure_score = 8 + min(12, structure_count * 2)
    section_scores = {
        "岗位匹配": min(100, jd_score * 5),
        "项目证据": min(100, (project_score + action_score) * 3),
        "量化结果": min(100, metric_score * 4),
        "工具链": min(100, tool_score * 4),
        "表达结构": min(100, (structure_score + length_score) * 3),
    }
    final_score = round(sum(section_scores.values()) / len(section_scores))
    risks = []
    actions = []
    brutal_comments = []
    evidence_gaps = []
    strengths = []
    if len(text) < 500:
        risks.append("简历内容偏短，HR 很难判断项目深度和个人贡献。")
        actions.append("补充 1-2 个项目经历，每个项目写清背景、职责、技术/工具、结果。")
        brutal_comments.append("目前更像“我学过什么”的说明，不像“我解决过什么问题”的简历。")
        evidence_gaps.append("缺少足够的项目上下文：业务背景、负责模块、输入输出和结果都需要补。")
    if not has_project:
        risks.append("项目经历证据不足，像技能清单而不是可验证经历。")
        actions.append("把技能放进具体项目动作，例如接口测试、自动化脚本、缺陷定位、性能验证。")
        brutal_comments.append("如果只写技能名，面试官很容易追问一句“你具体做了什么”，然后简历会失去说服力。")
        evidence_gaps.append("缺项目证据：至少写清一个项目的目标、模块、职责、难点、结果。")
    if not has_metrics:
        risks.append("缺少量化结果，表达可信度和竞争力不足。")
        actions.append("给项目补充数量、覆盖范围、缺陷数、接口数、响应时间或效率提升等指标。")
        brutal_comments.append("没有数字的项目经历会显得很虚，像课程作业介绍，不像可投递作品。")
        evidence_gaps.append("缺量化指标：用例数、接口数、缺陷数、响应时间、覆盖模块、优化前后对比。")
    if not matched_tools:
        risks.append("工具链表达不足，无法证明你真的做过开发/测试/联调。")
        actions.append("补充真实使用过的工具，例如 Flask、SQLite、Postman、JMeter、Pytest、Git 等。")
        evidence_gaps.append("缺工具链证据：技术栈和工具要跟项目动作绑定，不要单独堆在技能栏。")
    else:
        strengths.append("工具链有可用基础：" + "、".join(matched_tools[:6]))
    if action_count < 2:
        risks.append("个人贡献不够清楚，容易被理解成只是参与或旁观。")
        actions.append("多使用“负责、设计、实现、验证、定位、推动修复、输出报告”等动作词。")
        evidence_gaps.append("缺个人动作：每条项目经历都要能看出你本人负责了哪一段。")
    if result_count < 1:
        actions.append("每段经历最后补一句结果：解决了什么问题、带来什么改进、沉淀了什么产物。")
        evidence_gaps.append("缺结果闭环：没有体现修复、提升、覆盖、稳定性或交付产物。")
    if missing:
        actions.append("把 JD 缺口关键词补进项目经历语境：" + "、".join(missing[:6]))
        brutal_comments.append("这份简历和目标 JD 还有明显错位，关键词不能硬塞，要落到项目动作里。")
        evidence_gaps.append("JD 缺口关键词：" + "、".join(missing[:8]))
    if matched:
        strengths.append("JD 已命中关键词：" + "、".join(matched[:8]))
    if not risks:
        risks.append("整体基础可用，下一步重点是让项目贡献更具体、更贴目标岗位。")
    if not brutal_comments:
        brutal_comments.append("基础已经能看出方向，但还需要把“项目做了什么”升级成“你怎么判断、怎么实现、怎么验证”。")
    if not strengths:
        strengths.append("已有内容可以作为初稿，但需要补充项目证据和岗位关键词后再投递。")
    project_suggestions = [
        "项目描述按“背景-目标-职责-动作-结果”组织，避免只写系统有什么功能。",
        "测试岗位要突出测试用例设计、接口验证、缺陷闭环、自动化/性能工具。",
        "AI 项目要说明模型网关、兜底策略、提示词设计、业务流程和异常处理。",
        "把本系统写成可演示项目：多模型网关、简历/JD 分析、模拟面试状态流、语音表达评分、投递看板。",
        "每个亮点都配一条证据：接口、数据表、测试场景、异常处理或用户流程。"
    ]
    return {
        "score": max(35, min(96, final_score)),
        "section_scores": section_scores,
        "positioning": f"面向 {job_title or '目标岗位'} 的项目型候选人，核心卖点应落在真实项目、工具链、问题闭环和可验证结果。",
        "keywords": keywords,
        "matched_keywords": matched,
        "missing_keywords": missing,
        "strengths": strengths,
        "brutal_comments": brutal_comments,
        "evidence_gaps": evidence_gaps,
        "risks": risks,
        "actions": actions,
        "project_suggestions": project_suggestions,
    }


def build_improved_resume(resume_text: str, job_title: str, jd: str = "") -> dict:
    audit = build_resume_audit(resume_text, job_title, jd)
    local = tailor_resume_locally(resume_text, job_title, jd)
    ai_result = get_ai_client().optimize_resume(resume_text, job_title, jd) if get_ai_client().api_key else {"success": False, "content": ""}
    improved_resume = ai_result["content"] if ai_result.get("success") else local["tailored_resume"]
    strategy = [
        "保留真实经历，不编造公司和夸张结果。",
        "把“做过功能”改成“负责什么、如何验证、产出什么”。",
        "优先补齐 JD 高频词，并把关键词放进项目证据里。",
        "生成新版本而不是覆盖原简历，方便对比和回滚。",
    ]
    return {
        "audit": audit,
        "strategy": strategy,
        "improved_resume": improved_resume,
        "ai_used": bool(ai_result.get("success")),
    }


@app.route("/api/resumes/<int:resume_id>/optimize", methods=["POST"])
def optimize_resume(resume_id):
    row = get_resume_or_404(resume_id)
    if not row:
        return jsonify({"success": False, "message": "简历不存在"}), 404
    data = request.get_json() or {}
    result = get_ai_client().optimize_resume(row["content"], data.get("job_title", ""), data.get("jd", ""))
    return jsonify({"success": True, "suggestions": result["content"], "ai_used": result["success"], "provider": result["provider"]})


@app.route("/api/resumes/<int:resume_id>/tailor", methods=["POST"])
def tailor_resume(resume_id):
    row = get_resume_or_404(resume_id)
    if not row:
        return jsonify({"success": False, "message": "简历不存在"}), 404
    data = request.get_json() or {}
    job_title = (data.get("job_title") or "目标岗位").strip()
    jd = (data.get("jd") or data.get("job_requirements") or "").strip()
    tailored = tailor_resume_locally(row["content"], job_title, jd)
    ai_result = get_ai_client().optimize_resume(row["content"], job_title, jd) if jd else {"success": False, "content": ""}
    if ai_result.get("success"):
        tailored["ai_rewrite"] = ai_result["content"]
    with get_db() as conn:
        conn.execute("UPDATE resumes SET tailored_result = ? WHERE id = ?", (json.dumps(tailored, ensure_ascii=False), resume_id))
    return jsonify({"success": True, **tailored, "ai_used": bool(ai_result.get("success"))})


@app.route("/api/job-match", methods=["POST"])
def job_match():
    data = request.get_json() or {}
    resume_id = int(data.get("resume_id") or 0)
    row = get_resume_or_404(resume_id)
    if not row:
        return jsonify({"success": False, "message": "请先选择简历"}), 404
    job_title = data.get("job_title", "目标岗位")
    jd = data.get("job_requirements") or data.get("jd") or ""
    score, matched, missing = score_resume_against_jd(row["content"], jd + " " + job_title)
    ai = get_ai_client().match_job(row["content"], job_title, jd)
    analysis = ai["content"]
    with get_db() as conn:
        conn.execute(
            "INSERT INTO job_matches (user_id, resume_id, job_title, match_score, analysis) VALUES (?, ?, ?, ?, ?)",
            (row["user_id"], resume_id, job_title, score, analysis),
        )
    return jsonify({"success": True, "match_score": score, "matched_keywords": matched, "missing_keywords": missing, "analysis": analysis, "ai_used": ai["success"]})


@app.route("/api/interview/sessions", methods=["POST"])
def start_interview_session():
    data = request.get_json() or {}
    session_id = uuid.uuid4().hex[:12]
    resume_id = data.get("resume_id")
    resume = get_resume_or_404(int(resume_id)) if resume_id else None
    job_title = data.get("job_title", "目标岗位")
    session = {
        "session_id": session_id,
        "user_id": data.get("user_id", 1),
        "resume_id": resume_id,
        "job_title": job_title,
        "jd": data.get("jd", ""),
        "mode": data.get("mode", "standard"),
        "stage_index": 0,
        "conversation": [],
        "resume_content": resume["content"] if resume else "",
    }
    question = "欢迎参加模拟面试。请先做一个 2 分钟自我介绍，重点说清楚目标岗位、核心项目和你的优势。"
    session["conversation"].append({"role": "interviewer", "stage": "opening", "content": question})
    INTERVIEW_SESSIONS[session_id] = session
    return jsonify({"success": True, "session_id": session_id, "stage": "opening", "question": question, "progress": 1, "total": 6})


@app.route("/api/interview/sessions/<session_id>/answer", methods=["POST"])
def answer_interview_session(session_id):
    session = INTERVIEW_SESSIONS.get(session_id)
    if not session:
        return jsonify({"success": False, "message": "面试会话不存在"}), 404
    data = request.get_json() or {}
    answer = (data.get("answer") or "").strip()
    if not answer:
        return jsonify({"success": False, "message": "回答不能为空"}), 400
    answer_intent = detect_answer_intent(answer)
    voice = analyze_voice_text(answer, data.get("duration_seconds"))
    if answer_intent in {"skip", "too_short"}:
        voice["overall_score"] = 0
        voice["dimension_scores"] = {"表达流畅": 0, "结构逻辑": 0, "岗位相关": 0, "信息密度": 0}
    session["conversation"].append({"role": "candidate", "content": answer, "voice": voice})
    stages = [
        ("resume_deep_dive", "我看到你简历里有项目经历。请展开讲一个最能体现你能力的项目，按 STAR 结构回答。"),
        ("technical", f"如果你来测试/建设 {session['job_title']} 相关系统，你会如何设计核心用例和接口验证？"),
        ("behavioral", "讲一次你发现问题并推动解决的经历，你做了什么，结果怎样？"),
        ("candidate_questions", "现在进入反问环节。你会向面试官提哪两个问题？"),
        ("finished", "面试结束。系统已生成综合反馈。"),
    ]
    session["stage_index"] += 1
    stage, next_question = stages[min(session["stage_index"] - 1, len(stages) - 1)]
    if answer_intent in {"skip", "too_short"}:
        feedback = skipped_feedback(session["job_title"], stage)
    else:
        feedback = {
            "score": voice["overall_score"],
            "summary": build_interview_summary(answer, voice, session["job_title"]),
            "voice": voice,
            "suggestions": voice["tips"],
            "answer_upgrade": build_answer_upgrade(answer, session["job_title"]),
        }
    if stage == "finished":
        score = int(sum(item.get("voice", {}).get("overall_score", 75) for item in session["conversation"] if item["role"] == "candidate") / max(1, len([item for item in session["conversation"] if item["role"] == "candidate"])))
        feedback["score"] = score
        feedback["summary"] = "整体流程完成。建议把自我介绍压缩到 120 秒内，并准备 2 个项目深挖版本、1 个问题定位案例和 1 个团队协作案例。"
        with get_db() as conn:
            conn.execute(
                "INSERT INTO interviews (user_id, resume_id, job_title, conversation, score, feedback) VALUES (?, ?, ?, ?, ?, ?)",
                (session["user_id"], session.get("resume_id"), session["job_title"], json.dumps(session["conversation"], ensure_ascii=False), score, json.dumps(feedback, ensure_ascii=False)),
            )
    else:
        session["conversation"].append({"role": "interviewer", "stage": stage, "content": next_question})
    return jsonify({"success": True, "session_id": session_id, "stage": stage, "question": next_question, "feedback": feedback, "progress": min(session["stage_index"] + 1, 6), "total": 6})


def build_interview_summary(answer: str, voice: dict, job_title: str) -> str:
    if len(answer) < 60:
        return f"回答偏短。面试 {job_title} 时，需要把“做过什么、怎么做、结果如何”讲完整。"
    if voice["structure_score"] < 2:
        return "内容有素材，但结构不够明显。建议先给结论，再按背景、行动、结果展开。"
    if voice["filler_count"] > 2:
        return "信息量可以，但口头禅偏多。建议用短暂停顿替代“然后、就是、那个”。"
    return "回答整体可用，已经具备项目证据。下一步重点补充量化指标和岗位关键词。"


def build_answer_upgrade(answer: str, job_title: str) -> str:
    return (
        f"可升级表达：面向 {job_title}，我在项目中不仅参与实现/测试，还围绕核心业务流程设计验证方案。"
        "例如在 AI 求职辅助系统中，我覆盖了简历上传、JD 匹配、模拟面试和投递看板等流程，"
        "通过接口测试、异常场景和回归验证保证系统稳定，并把测试结论沉淀成报告。"
    )


def detect_answer_intent(answer: str) -> str:
    text = re.sub(r"\s+", "", (answer or "").lower())
    if not text:
        return "empty"
    skip_words = ["不知道", "不会", "不清楚", "没想好", "下一题", "跳过", "pass", "next", "不会答"]
    if any(word in text for word in skip_words):
        return "skip"
    if len(text) < 8:
        return "too_short"
    return "answer"


def skipped_feedback(job_title: str, stage_name_text: str = "本题") -> dict:
    voice = analyze_voice_text("不知道")
    voice["overall_score"] = 0
    voice["dimension_scores"] = {"表达流畅": 0, "结构逻辑": 0, "岗位相关": 0, "信息密度": 0}
    return {
        "score": 0,
        "summary": f"你选择跳过{stage_name_text}。这在练习里可以，但真实面试不能只说不知道。建议先给一个诚实回应，再说你的补救思路。",
        "voice": voice,
        "suggestions": [
            "可用话术：这个点我现在不能完整回答，但我会先确认概念，再结合项目场景补充验证。",
            "遇到不会的题，至少说出你知道的边界、排查路径或学习计划。",
            "系统已进入下一题/下一阶段，不会把跳过内容包装成虚假能力。"
        ],
        "answer_upgrade": f"保底回答：这个问题我还需要补充学习。面向 {job_title}，我会从岗位要求出发，先查清概念，再用项目里的真实场景做验证和复盘。",
    }


@app.route("/api/interview/analyze-voice", methods=["POST"])
def analyze_voice_answer():
    data = request.get_json() or {}
    answer = data.get("answer", "")
    if not answer:
        return jsonify({"success": False, "message": "请先输入或录入回答内容"}), 400
    return jsonify({"success": True, **analyze_voice_text(answer, data.get("duration_seconds"), data.get("audio_metrics"))})


@app.route("/api/interview/analyze-audio", methods=["POST"])
def analyze_audio_answer():
    transcript = (request.form.get("transcript") or "").strip()
    if not transcript:
        return jsonify({"success": False, "message": "请提供录音对应的转写文本"}), 400
    try:
        metrics = json.loads(request.form.get("metrics") or "{}")
    except json.JSONDecodeError:
        metrics = {}
    audio_file = request.files.get("audio")
    saved_name = ""
    if audio_file and audio_file.filename:
        saved_name = safe_filename(f"audio_{uuid.uuid4().hex[:8]}_{audio_file.filename}")
        audio_file.save(os.path.join(UPLOAD_FOLDER, saved_name))
    result = analyze_voice_text(transcript, request.form.get("duration_seconds"), metrics)
    result["summary"] = build_audio_summary(result, saved_name)
    with get_db() as conn:
        conn.execute(
            "INSERT INTO audio_records (user_id, transcript, audio_file, score, metrics, feedback) VALUES (?, ?, ?, ?, ?, ?)",
            (
                int(request.form.get("user_id") or 1),
                transcript,
                saved_name,
                result.get("overall_score"),
                json.dumps(metrics, ensure_ascii=False),
                json.dumps(result, ensure_ascii=False),
            ),
        )
    return jsonify({"success": True, "audio_file": saved_name, **result})


def build_audio_summary(result: dict, saved_name: str = "") -> str:
    metrics = result.get("audio_metrics") or {}
    file_note = "已保存录音，可用于复盘。" if saved_name else "未保存音频文件，仅使用浏览器侧音频指标。"
    return (
        f"{file_note} 本次语速为 {result['estimated_speech_rate']} 字/分钟，"
        f"停顿比例约 {round(float(metrics.get('silence_ratio') or 0) * 100)}%，"
        f"口头禅 {result['filler_count']} 次。建议重点关注：{result['audio_quality']}"
    )


@app.route("/api/interview/professional-pack", methods=["POST"])
def professional_pack():
    data = request.get_json() or {}
    category = data.get("category", "test")
    level = data.get("level", "campus")
    job_title = data.get("job_title", "目标岗位")
    bank = question_bank()
    base_questions = bank.get(category, bank["general"])
    questions = []
    level_name = {"campus": "校招基础", "junior": "初级实战", "project": "项目深挖"}.get(level, "校招基础")
    for item in base_questions[:5]:
        questions.append({
            "question": item["question"],
            "reference": item["answer"],
            "focus": f"{job_title} · {category}",
            "difficulty": level_name,
        })
    questions.extend(build_project_followup_questions(category, job_title, level))
    return jsonify({"success": True, "category": category, "level": level, "questions": questions[:8]})


def build_project_followup_questions(category: str, job_title: str, level: str) -> list[dict]:
    if category == "test":
        return [
            {
                "question": f"如果让你测试一个 {job_title} 相关的 AI Web 系统，你会如何拆分测试范围？",
                "reference": "可以从核心业务流程、接口契约、文件上传解析、模型异常兜底、权限和兼容性几个维度拆分，并说明优先级。",
                "focus": "测试设计能力",
                "difficulty": "项目场景",
            },
            {
                "question": "模型接口偶发超时或返回格式不稳定，你会怎么设计验证和兜底方案？",
                "reference": "回答要覆盖超时重试、错误提示、本地规则兜底、日志记录、接口 Mock、边界用例和回归验证。",
                "focus": "异常处理",
                "difficulty": "实战追问",
            },
        ]
    if category == "ai":
        return [
            {
                "question": "你怎么理解 AI Agent 和普通聊天机器人的区别？结合本项目说明。",
                "reference": "Agent 应有目标、状态、工具/数据调用和任务推进。本项目可举简历、JD、面试、投递数据联动的例子。",
                "focus": "Agent 理解",
                "difficulty": "项目深挖",
            },
            {
                "question": "多模型网关为什么需要本地兜底？如何判断兜底结果是否可用？",
                "reference": "从演示稳定性、成本、网络/API 异常说起，再讲规则模板、关键词抽取、单元测试和人工可解释性。",
                "focus": "模型工程",
                "difficulty": "实战追问",
            },
        ]
    if category == "python":
        return [
            {
                "question": "Flask 项目里你会如何划分路由、服务逻辑和数据访问？",
                "reference": "说明接口层负责参数和响应，服务层负责业务，数据层负责 SQL；再补充错误处理和测试策略。",
                "focus": "后端结构",
                "difficulty": "基础到实战",
            }
        ]
    return [
        {
            "question": "前端页面如何保证复杂表单、结果区和模块切换不混乱？",
            "reference": "可以从状态管理、模块过滤、输入校验、结果复用、响应式布局和用户下一步引导回答。",
            "focus": "前端工程",
            "difficulty": "项目场景",
        }
    ]


@app.route("/api/interview/practice-feedback", methods=["POST"])
def practice_feedback():
    data = request.get_json() or {}
    question = (data.get("question") or "").strip()
    answer = (data.get("answer") or "").strip()
    category = data.get("category", "general")
    if not question or not answer:
        return jsonify({"success": False, "message": "题目和回答不能为空"}), 400

    if detect_answer_intent(answer) in {"skip", "too_short"}:
        result = {
            "success": True,
            "score": 0,
            "category": category,
            "question": question,
            "dimension_scores": {"专业性": 0, "结构化": 0, "完整度": 0, "表达": 0},
            "hits": [],
            "problems": ["本题没有形成有效回答，系统不会编造评分。"],
            "sample_answer": build_sample_practice_answer(question, category),
            "upgrade": "不会的问题建议诚实说明边界，再补一个排查思路或学习计划。",
            "follow_up": "下一步：先看参考答案，再用自己的项目经历重答一遍。",
            "needs_answer": True,
        }
        save_practice_record(data.get("user_id", 1), category, question, answer, result)
        return jsonify(result)

    voice = analyze_voice_text(answer)
    technical_terms = extract_keywords(answer)
    structure_hit = any(word in answer for word in ["首先", "其次", "最后", "背景", "任务", "行动", "结果", "因此"])
    score = voice["overall_score"]
    if category in {"test", "python", "frontend", "ai"}:
        score = min(96, score + min(12, len(technical_terms) * 3))
    if not structure_hit:
        score = max(35, score - 8)

    sample_answer = build_sample_practice_answer(question, category)
    result = {
        "success": True,
        "score": score,
        "category": category,
        "question": question,
        "dimension_scores": {
            "专业性": min(95, 48 + len(technical_terms) * 8),
            "结构化": 86 if structure_hit else 55,
            "完整度": min(95, 40 + len(answer) // 3),
            "表达": voice["dimension_scores"]["表达流畅"],
        },
        "hits": technical_terms,
        "problems": [
            item for item in [
                None if structure_hit else "回答缺少清晰结构，建议使用“结论-步骤-结果”。",
                None if len(answer) >= 80 else "回答偏短，需要补充例子或项目经历。",
                None if technical_terms else "专业关键词较少，建议加入工具、方法或指标。",
            ] if item
        ],
        "sample_answer": sample_answer,
        "upgrade": build_answer_upgrade(answer, data.get("job_title", "目标岗位")),
        "follow_up": build_follow_up_question(question, category),
    }
    save_practice_record(data.get("user_id", 1), category, question, answer, result)
    return jsonify(result)


def save_practice_record(user_id: int, category: str, question: str, answer: str, result: dict) -> None:
    with get_db() as conn:
        conn.execute(
            "INSERT INTO practice_records (user_id, category, question, answer, score, feedback) VALUES (?, ?, ?, ?, ?, ?)",
            (user_id, category, question, answer, result.get("score"), json.dumps(result, ensure_ascii=False)),
        )


def build_follow_up_question(question: str, category: str) -> str:
    if category == "test":
        return "追问：如果这个功能上线后出现偶发失败，你如何定位是前端、后端、数据库还是模型接口问题？"
    if category == "ai":
        return "追问：如果模型返回不符合预期，你会如何通过 prompt、规则校验和兜底策略保证产品可用？"
    if category == "python":
        return "追问：这个接口如果并发访问变多，你会从哪些指标判断瓶颈？"
    if category == "frontend":
        return "追问：移动端和桌面端布局差异较大时，你如何做断点和视觉回归验证？"
    return "追问：请把这个回答再结合一个真实项目经历讲一遍，重点说你的个人贡献。"


def build_sample_practice_answer(question: str, category: str) -> str:
    if category == "test":
        return "参考回答：我会先确认需求和核心业务流程，再从正常流程、边界值、异常输入、权限、接口契约、兼容性和性能几个维度设计用例。执行时会记录实际结果、缺陷复现步骤和优先级，最后通过回归测试确认问题闭环。"
    if category == "ai":
        return "参考回答：AI Agent 不只是单轮聊天，它需要明确目标、保存上下文、按流程推进任务，并在模型不可用时有兜底策略。比如本项目把简历、JD、面试状态和语音分析串起来，形成完整求职辅助流程。"
    if category == "python":
        return "参考回答：Flask 项目可以按路由、服务逻辑、数据访问和配置拆分。接口层负责参数校验和响应，服务层处理业务规则，数据库层负责持久化，同时统一错误处理，方便测试和维护。"
    return "参考回答：我会先给结论，再用一个真实项目举例说明背景、我的行动和结果，最后回到岗位要求，说明这段经历为什么能证明我适合这个岗位。"


@app.route("/api/training-records/<int:user_id>")
def list_training_records(user_id):
    with get_db() as conn:
        interviews = conn.execute(
            "SELECT id, job_title, conversation, score, feedback, created_at FROM interviews WHERE user_id = ? ORDER BY created_at DESC LIMIT 30",
            (user_id,),
        ).fetchall()
        practices = conn.execute(
            "SELECT id, category, question, answer, score, feedback, created_at FROM practice_records WHERE user_id = ? ORDER BY created_at DESC LIMIT 50",
            (user_id,),
        ).fetchall()
        audios = conn.execute(
            "SELECT id, transcript, audio_file, score, metrics, feedback, created_at FROM audio_records WHERE user_id = ? ORDER BY created_at DESC LIMIT 50",
            (user_id,),
        ).fetchall()
    return jsonify({
        "success": True,
        "interviews": [dict(row) for row in interviews],
        "practices": [dict(row) for row in practices],
        "audios": [dict(row) for row in audios],
    })


@app.route("/api/training-records/<record_type>/<int:record_id>", methods=["DELETE"])
def delete_training_record(record_type, record_id):
    table_map = {
        "interview": "interviews",
        "practice": "practice_records",
        "audio": "audio_records",
    }
    table = table_map.get(record_type)
    if not table:
        return jsonify({"success": False, "message": "记录类型不存在"}), 400
    with get_db() as conn:
        row = conn.execute(f"SELECT * FROM {table} WHERE id = ?", (record_id,)).fetchone()
        if not row:
            return jsonify({"success": False, "message": "记录不存在"}), 404
        if record_type == "audio" and row["audio_file"]:
            audio_path = os.path.join(UPLOAD_FOLDER, row["audio_file"])
            if os.path.exists(audio_path):
                try:
                    os.remove(audio_path)
                except OSError:
                    pass
        conn.execute(f"DELETE FROM {table} WHERE id = ?", (record_id,))
    return jsonify({"success": True, "message": "记录已删除"})


@app.route("/api/training-records/<int:user_id>/clear", methods=["DELETE"])
def clear_training_records(user_id):
    with get_db() as conn:
        audio_rows = conn.execute("SELECT audio_file FROM audio_records WHERE user_id = ?", (user_id,)).fetchall()
        for row in audio_rows:
            if row["audio_file"]:
                audio_path = os.path.join(UPLOAD_FOLDER, row["audio_file"])
                if os.path.exists(audio_path):
                    try:
                        os.remove(audio_path)
                    except OSError:
                        pass
        conn.execute("DELETE FROM interviews WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM practice_records WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM audio_records WHERE user_id = ?", (user_id,))
    return jsonify({"success": True, "message": "训练记录已清空"})


@app.route("/api/agent/chat", methods=["POST"])
def agent_chat():
    data = request.get_json() or {}
    message = data.get("message", "")
    if not message:
        return jsonify({"success": False, "message": "请输入问题"}), 400
    context = data.get("context") or build_agent_runtime_context(data.get("user_id", 1))
    result = get_ai_client().agent_chat(message, context)
    return jsonify({"success": True, "reply": result["content"], "ai_used": result["success"], "provider": result["provider"]})


@app.route("/api/career/report/<int:user_id>", methods=["POST"])
def career_report(user_id):
    with get_db() as conn:
        resumes = conn.execute("SELECT title, content, updated_at FROM resumes WHERE user_id = ? ORDER BY updated_at DESC LIMIT 3", (user_id,)).fetchall()
        matches = conn.execute("SELECT job_title, match_score, created_at FROM job_matches WHERE user_id = ? ORDER BY created_at DESC LIMIT 5", (user_id,)).fetchall()
        interviews = conn.execute("SELECT job_title, score, feedback, created_at FROM interviews WHERE user_id = ? ORDER BY created_at DESC LIMIT 5", (user_id,)).fetchall()
        apps = conn.execute("SELECT company, job_title, status, city, notes FROM job_applications WHERE user_id = ? ORDER BY updated_at DESC LIMIT 8", (user_id,)).fetchall()
    report = build_career_report(resumes, matches, interviews, apps)
    if get_ai_client().api_key:
        result = get_ai_client().chat([
            {"role": "system", "content": "你是求职策略教练。请把用户当前求职数据整理成一份结构清晰、行动明确的中文作战报告。"},
            {"role": "user", "content": report},
        ], temperature=0.35, max_tokens=1100)
        if result.get("success"):
            report = result["content"]
    return jsonify({"success": True, "report": report})


def build_career_report(resumes: list[sqlite3.Row], matches: list[sqlite3.Row], interviews: list[sqlite3.Row], apps: list[sqlite3.Row]) -> str:
    resume_line = "暂无简历，第一优先级是录入一份可分析简历。"
    if resumes:
        resume_line = "已保存简历：" + "、".join(row["title"] for row in resumes)
    match_line = "暂无 JD 匹配记录，建议先选 1 个真实岗位做定制优化。"
    if matches:
        avg = round(sum(row["match_score"] or 0 for row in matches) / len(matches))
        match_line = f"最近 {len(matches)} 次 JD 匹配平均分约 {avg}，重点看低分岗位的关键词缺口。"
    interview_line = "暂无模拟面试记录，建议先跑一轮完整流程。"
    if interviews:
        avg = round(sum(row["score"] or 0 for row in interviews) / len(interviews))
        interview_line = f"最近 {len(interviews)} 次面试训练平均分约 {avg}，建议继续打磨自我介绍和项目深挖。"
    app_line = "暂无投递记录，建议建立投递看板，避免只投不跟进。"
    if apps:
        status_count = {}
        for row in apps:
            status_count[row["status"]] = status_count.get(row["status"], 0) + 1
        app_line = "当前投递阶段分布：" + "、".join(f"{key}{value}条" for key, value in status_count.items())
    return (
        "## 求职作战报告\n"
        f"### 1. 简历资产\n{resume_line}\n\n"
        f"### 2. 岗位匹配\n{match_line}\n\n"
        f"### 3. 面试训练\n{interview_line}\n\n"
        f"### 4. 投递推进\n{app_line}\n\n"
        "### 5. 下一步建议\n"
        "- 先选一个真实 JD 做简历定制，生成匹配分和缺口清单。\n"
        "- 把 JD 自动带入模拟面试，完成一轮自我介绍、项目深挖、技术追问和反问。\n"
        "- 将目标公司加入投递看板，按阶段推进并生成跟进话术。\n"
        "- 把最终优化后的项目经历导出为 PDF/Word，用于真实投递和实训展示。"
    )


def build_agent_runtime_context(user_id: int = 1) -> str:
    """Give the coach enough product context to answer like a job-search agent."""
    with get_db() as conn:
        resume_count = conn.execute("SELECT COUNT(*) FROM resumes WHERE user_id = ?", (user_id,)).fetchone()[0]
        latest_resume = conn.execute(
            "SELECT title, content FROM resumes WHERE user_id = ? ORDER BY updated_at DESC LIMIT 1",
            (user_id,),
        ).fetchone()
        apps = conn.execute(
            "SELECT company, job_title, status, city FROM job_applications WHERE user_id = ? ORDER BY updated_at DESC LIMIT 5",
            (user_id,),
        ).fetchall()
        interviews = conn.execute(
            "SELECT job_title, score FROM interviews WHERE user_id = ? ORDER BY created_at DESC LIMIT 3",
            (user_id,),
        ).fetchall()
    latest_resume_text = ""
    if latest_resume:
        latest_resume_text = f"最近简历：{latest_resume['title']}\n简历摘要：{latest_resume['content'][:900]}"
    app_text = "\n".join(f"- {row['company']} / {row['job_title']} / {row['status']} / {row['city'] or '未填城市'}" for row in apps)
    interview_text = "\n".join(f"- {row['job_title']}：{row['score']}分" for row in interviews)
    return (
        "你是职途AI里的求职智能体教练。回答要具体、可执行、贴近应届生/实训项目包装场景，"
        "优先给简历改写、JD匹配、模拟面试、投递追踪、谈薪准备的下一步动作。\n"
        f"当前用户简历数量：{resume_count}\n"
        f"{latest_resume_text or '暂无可用简历内容。'}\n"
        f"最近投递：\n{app_text or '暂无投递记录。'}\n"
        f"最近面试训练：\n{interview_text or '暂无面试训练记录。'}"
    )


@app.route("/api/questions")
def questions():
    category = request.args.get("category", "all")
    bank = question_bank()
    if category == "all":
        data = [{"category": key, **item} for key, items in bank.items() for item in items]
    else:
        data = [{"category": category, **item} for item in bank.get(category, [])]
    return jsonify({"success": True, "data": data, "categories": list(bank.keys())})


@app.route("/api/applications", methods=["POST"])
def create_application():
    data = request.get_json() or {}
    with get_db() as conn:
        cursor = conn.execute(
            "INSERT INTO job_applications (user_id, company, job_title, status, city, salary_min, salary_max, notes) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (
                data.get("user_id", 1),
                data.get("company", "未命名公司"),
                data.get("job_title", "目标岗位"),
                data.get("status", "已投递"),
                data.get("city", ""),
                data.get("salary_min"),
                data.get("salary_max"),
                data.get("notes", ""),
            ),
        )
    return jsonify({"success": True, "application_id": cursor.lastrowid})


@app.route("/api/applications/<int:user_id>")
def list_applications(user_id):
    with get_db() as conn:
        rows = conn.execute("SELECT * FROM job_applications WHERE user_id = ? ORDER BY updated_at DESC", (user_id,)).fetchall()
    return jsonify({"success": True, "data": [dict(row) for row in rows]})


@app.route("/api/applications/detail/<int:application_id>")
def application_detail(application_id):
    with get_db() as conn:
        row = conn.execute("SELECT * FROM job_applications WHERE id = ?", (application_id,)).fetchone()
    if not row:
        return jsonify({"success": False, "message": "投递记录不存在"}), 404
    return jsonify({"success": True, "data": dict(row)})


@app.route("/api/applications/<int:application_id>", methods=["PUT"])
def update_application(application_id):
    data = request.get_json() or {}
    with get_db() as conn:
        cursor = conn.execute(
            """
            UPDATE job_applications
            SET company = ?, job_title = ?, status = ?, city = ?, notes = ?, updated_at = CURRENT_TIMESTAMP
            WHERE id = ?
            """,
            (
                data.get("company", "未命名公司"),
                data.get("job_title", "目标岗位"),
                data.get("status", "已投递"),
                data.get("city", ""),
                data.get("notes", ""),
                application_id,
            ),
        )
    if cursor.rowcount == 0:
        return jsonify({"success": False, "message": "投递记录不存在"}), 404
    return jsonify({"success": True, "message": "投递记录已更新"})


@app.route("/api/applications/<int:application_id>", methods=["DELETE"])
def delete_application(application_id):
    with get_db() as conn:
        cursor = conn.execute("DELETE FROM job_applications WHERE id = ?", (application_id,))
    if cursor.rowcount == 0:
        return jsonify({"success": False, "message": "投递记录不存在"}), 404
    return jsonify({"success": True, "message": "投递记录已删除"})


@app.route("/api/applications/<int:application_id>/coach", methods=["POST"])
def coach_application(application_id):
    data = request.get_json() or {}
    user_id = int(data.get("user_id", 1))
    with get_db() as conn:
        app_row = conn.execute("SELECT * FROM job_applications WHERE id = ?", (application_id,)).fetchone()
        resume_row = conn.execute(
            "SELECT title, content FROM resumes WHERE user_id = ? ORDER BY updated_at DESC LIMIT 1",
            (user_id,),
        ).fetchone()
        interview_row = conn.execute(
            "SELECT job_title, score, feedback FROM interviews WHERE user_id = ? ORDER BY created_at DESC LIMIT 1",
            (user_id,),
        ).fetchone()
    if not app_row:
        return jsonify({"success": False, "message": "投递记录不存在"}), 404

    plan = build_application_followup_plan(app_row, resume_row, interview_row)
    ai_note = ""
    if get_ai_client().api_key:
        result = get_ai_client().chat([
            {"role": "system", "content": "你是求职投递教练。请根据投递阶段给出简洁、可执行的跟进建议，不要空话。"},
            {"role": "user", "content": json.dumps({
                "application": dict(app_row),
                "latest_resume": dict(resume_row) if resume_row else None,
                "latest_interview": dict(interview_row) if interview_row else None,
                "local_plan": plan,
            }, ensure_ascii=False)},
        ], temperature=0.35, max_tokens=700)
        ai_note = result.get("content", "") if result.get("success") else ""
    return jsonify({"success": True, **plan, "ai_note": ai_note})


@app.route("/api/applications/<int:application_id>/advance", methods=["POST"])
def advance_application(application_id):
    stages = ["已投递", "简历筛选", "笔试", "一面", "二面", "HR 面", "Offer", "已拒绝"]
    with get_db() as conn:
        row = conn.execute("SELECT * FROM job_applications WHERE id = ?", (application_id,)).fetchone()
        if not row:
            return jsonify({"success": False, "message": "投递记录不存在"}), 404
        current = row["status"]
        try:
            next_status = stages[min(stages.index(current) + 1, len(stages) - 1)]
        except ValueError:
            next_status = "简历筛选"
        conn.execute(
            "UPDATE job_applications SET status = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            (next_status, application_id),
        )
    return jsonify({"success": True, "status": next_status})


def build_application_followup_plan(app_row: sqlite3.Row, resume_row: sqlite3.Row | None, interview_row: sqlite3.Row | None) -> dict:
    company = app_row["company"]
    job = app_row["job_title"]
    status = app_row["status"]
    city = app_row["city"] or "目标城市"
    notes = app_row["notes"] or ""
    resume_hint = f"最近简历《{resume_row['title']}》" if resume_row else "当前暂无已保存简历"
    interview_hint = ""
    if interview_row:
        interview_hint = f"最近一次 {interview_row['job_title']} 模拟面试得分 {interview_row['score']}，建议把低分维度转成复盘待办。"

    stage_rules = {
        "已投递": (
            "投递后 2-3 天检查岗位状态；如果没有反馈，补充一次礼貌跟进，并同步准备该岗位的 2 分钟项目介绍。",
            "风险是只投不跟进，后续忘记岗位要求；建议在备注里补充投递渠道、JD 关键词和截止时间。",
        ),
        "简历筛选": (
            "围绕 JD 关键词再做一次简历定制，准备一版更贴合该岗位的项目经历表述。",
            "风险是简历泛化，HR 看不到岗位相关证据；需要把工具、动作、结果写进项目经历。",
        ),
        "笔试": (
            "整理笔试范围，优先准备基础题、接口测试、SQL、Python/前端基础和项目场景题。",
            "风险是只刷题不复盘；建议记录错题类型，并把知识点转成面试可讲案例。",
        ),
        "一面": (
            "准备自我介绍、项目深挖、技术追问和反问问题；用模拟面试模块跑一轮完整流程。",
            "风险是项目讲得像流水账；建议用 STAR 结构讲清背景、动作、结果和个人贡献。",
        ),
        "二面": (
            "强化项目决策、问题定位、协作推动和结果量化，准备 2 个能体现成长性的案例。",
            "风险是回答停留在功能层；需要补充为什么这样设计、如何验证、遇到问题怎么取舍。",
        ),
        "HR 面": (
            "准备求职动机、稳定性、薪资预期和到岗时间；提前查城市和岗位薪资区间。",
            "风险是薪资表达没依据；建议用城市、经验、技能匹配度和 Offer 进度作为谈薪支撑。",
        ),
        "Offer": (
            "核对薪资结构、试用期、五险一金、工作地点、入职材料和违约条款，再做最终选择。",
            "风险是只看月薪不看总包和试用期规则；建议整理对比表后再确认。",
        ),
        "已拒绝": (
            "记录拒绝原因，把它转成简历、面试或技能的下一轮改进项。",
            "风险是只记录失败结果，不沉淀原因；建议复盘筛选、笔试、面试各阶段卡点。",
        ),
    }
    next_action, risk = stage_rules.get(status, stage_rules["已投递"])
    if notes:
        risk += f" 当前备注里已有线索：{notes[:120]}"
    template = (
        f"您好，我是投递 {company}「{job}」岗位的候选人。想礼貌确认一下目前流程进展。"
        f"我这边可以结合岗位要求补充 {resume_hint} 中与岗位更相关的项目材料，也可以配合后续笔试/面试安排。谢谢。"
    )
    if status in {"一面", "二面", "HR 面"}:
        template = (
            f"您好，感谢之前关于 {company}「{job}」岗位的沟通。"
            f"我已根据面试反馈继续梳理项目经历和岗位匹配点，想确认一下后续流程安排。谢谢。"
        )
    return {
        "title": f"{company} / {job} 跟进策略",
        "company": company,
        "job_title": job,
        "status": status,
        "city": city,
        "next_action": next_action,
        "risk": risk,
        "message_template": template,
        "resume_hint": resume_hint,
        "interview_hint": interview_hint,
    }


@app.route("/api/dashboard/<int:user_id>")
def dashboard(user_id):
    with get_db() as conn:
        resume_count = conn.execute("SELECT COUNT(*) FROM resumes WHERE user_id = ?", (user_id,)).fetchone()[0]
        interview_rows = conn.execute("SELECT score, created_at FROM interviews WHERE user_id = ? ORDER BY created_at", (user_id,)).fetchall()
        match_rows = conn.execute("SELECT match_score, created_at FROM job_matches WHERE user_id = ? ORDER BY created_at", (user_id,)).fetchall()
        app_rows = conn.execute("SELECT status, company, job_title, updated_at FROM job_applications WHERE user_id = ? ORDER BY updated_at DESC", (user_id,)).fetchall()
        practice_count = conn.execute("SELECT COUNT(*) FROM practice_records WHERE user_id = ?", (user_id,)).fetchone()[0]
        audio_count = conn.execute("SELECT COUNT(*) FROM audio_records WHERE user_id = ?", (user_id,)).fetchone()[0]
    stats = {
        "resumes": resume_count,
        "interviews": len(interview_rows),
        "matches": len(match_rows),
        "applications": len(app_rows),
        "practices": practice_count,
        "audios": audio_count,
    }
    return jsonify({
        "success": True,
        "stats": stats,
        "interview_scores": [dict(row) for row in interview_rows],
        "match_scores": [dict(row) for row in match_rows],
        "activities": [dict(row) for row in app_rows[:6]],
        "next_actions": build_next_actions(stats),
        "career_pulse": build_career_pulse(stats, interview_rows, match_rows, app_rows),
    })


def build_career_pulse(stats: dict, interview_rows: list[sqlite3.Row], match_rows: list[sqlite3.Row], app_rows: list[sqlite3.Row]) -> dict:
    readiness = 10
    blockers = []
    weekly_plan = []
    if stats["resumes"]:
        readiness += 22
    else:
        blockers.append("还没有可分析简历，后面的 JD 匹配和面试追问都缺上下文。")
        weekly_plan.append({"title": "录入一份主简历", "page": "resume", "module": "input"})
    if stats["matches"]:
        avg_match = round(sum(row["match_score"] or 0 for row in match_rows) / max(1, len(match_rows)))
        readiness += min(22, max(8, avg_match // 4))
        if avg_match < 65:
            blockers.append("最近 JD 匹配分偏低，先补岗位关键词和项目证据再投。")
    else:
        blockers.append("还没有用真实 JD 做过匹配，简历是否贴岗位还不清楚。")
        weekly_plan.append({"title": "复制一个真实 JD 做匹配", "page": "resume", "module": "jd"})
    if stats["interviews"]:
        scores = [row["score"] or 0 for row in interview_rows]
        avg_interview = round(sum(scores) / max(1, len(scores)))
        readiness += min(24, max(8, avg_interview // 4))
        if avg_interview < 60:
            blockers.append("面试训练平均分偏低，优先打磨自我介绍和项目深挖。")
    else:
        blockers.append("还没跑完整模拟面试，真实面试流程感不足。")
        weekly_plan.append({"title": "完成一轮模拟面试", "page": "interview", "module": "mock"})
    if stats["practices"]:
        readiness += min(12, 4 + stats["practices"] * 2)
    else:
        weekly_plan.append({"title": "做 3 道题库练习", "page": "interview", "module": "practice"})
    if stats["audios"]:
        readiness += 8
    else:
        blockers.append("还没有真实录音复盘，语速、停顿和口头禅没有被校准。")
    if stats["applications"]:
        readiness += min(12, 5 + stats["applications"] * 2)
    else:
        weekly_plan.append({"title": "建立投递看板", "page": "tracker", "module": "add"})
    status_count = {}
    for row in app_rows:
        status_count[row["status"]] = status_count.get(row["status"], 0) + 1
    readiness = max(0, min(100, readiness))
    label = "可投递" if readiness >= 72 else "需要打磨" if readiness >= 45 else "先补基础"
    if not weekly_plan:
        weekly_plan = [
            {"title": "复盘最近一次低分记录", "page": "interview", "module": "records"},
            {"title": "推进一个投递阶段", "page": "tracker", "module": "board"},
            {"title": "导出最终简历版本", "page": "resume", "module": "export"},
        ]
    return {
        "score": readiness,
        "label": label,
        "summary": f"当前准备度 {readiness}/100。系统按简历资产、JD 匹配、面试训练、语音复盘和投递推进综合评估。",
        "blockers": blockers[:4] or ["基础链路已经跑通，下一步重点是把真实 JD、真实录音和投递反馈持续沉淀。"],
        "weekly_plan": weekly_plan[:4],
        "funnel": status_count,
    }


def build_next_actions(stats: dict) -> list[dict]:
    actions = []
    if stats["resumes"] == 0:
        actions.append({
            "title": "先建立一份可分析的简历",
            "description": "上传 Word/PDF 或粘贴文本，系统才能做诊断、JD 匹配、导出和面试追问。",
            "page": "resume",
            "module": "input",
            "cta": "录入简历",
        })
    else:
        actions.append({
            "title": "用 JD 检查简历是否命中岗位",
            "description": "把目标岗位 JD 粘进去，系统会给出匹配分、关键词缺口和可讲述的项目亮点。",
            "page": "resume",
            "module": "jd",
            "cta": "做 JD 优化",
        })
    if stats["interviews"] == 0:
        actions.append({
            "title": "跑一轮完整模拟面试",
            "description": "从自我介绍、项目深挖到反问总结，训练结果会沉淀到 AI 教练上下文。",
            "page": "interview",
            "module": "mock",
            "cta": "开始面试",
        })
    if stats["applications"] == 0:
        actions.append({
            "title": "建立投递看板",
            "description": "记录公司、岗位、阶段和备注，后续可以生成跟进话术和谈薪准备。",
            "page": "tracker",
            "module": "add",
            "cta": "新增投递",
        })
    else:
        actions.append({
            "title": "推进投递状态并复盘反馈",
            "description": "把投递从已投递推进到笔试/面试/Offer，系统会按阶段给跟进建议。",
            "page": "tracker",
            "module": "board",
            "cta": "看投递板",
        })
    return actions[:3]


@app.route("/api/salary/evaluate", methods=["POST"])
def salary_evaluate():
    data = request.get_json() or {}
    city_factor = {"北京": 1.25, "上海": 1.25, "深圳": 1.2, "广州": 1.05, "杭州": 1.15, "成都": 0.9, "武汉": 0.85}.get(data.get("city"), 1)
    exp = data.get("experience", "应届生")
    base = {"应届生": 9000, "1-3年": 15000, "3-5年": 24000, "5年以上": 36000}.get(exp, 12000)
    skills_bonus = min(5000, int(data.get("skills_count") or 0) * 500)
    avg = int((base + skills_bonus) * city_factor)
    return jsonify({
        "success": True,
        "range": {"min": int(avg * 0.75), "avg": avg, "max": int(avg * 1.35)},
        "advice": "谈薪时用岗位 JD 匹配度、项目结果和可立即上手的工具链作为依据。",
    })


@app.route("/api/skills/radar", methods=["POST"])
def skills_radar():
    data = request.get_json() or {}
    text = data.get("resume_content", "")
    if data.get("resume_id"):
        row = get_resume_or_404(int(data["resume_id"]))
        text = row["content"] if row else text
    categories = {
        "编程基础": ["Python", "Java", "Flask", "Spring"],
        "测试能力": ["Selenium", "Pytest", "JMeter", "Postman", "接口测试", "自动化测试"],
        "工程工具": ["Git", "Docker", "Linux", "MySQL", "Redis"],
        "AI 应用": ["AI", "智能体", "大模型"],
        "表达呈现": ["报告", "文档", "沟通", "项目"],
    }
    radar = []
    for name, words in categories.items():
        matched = [word for word in words if word.lower() in text.lower()]
        score = min(10, 3 + len(matched) * 2)
        missing = [word for word in words if word not in matched][:4]
        suggestion = "继续补充项目证据，避免只在技能栏堆关键词。"
        if score <= 5:
            suggestion = f"建议补充 {name} 证据：{', '.join(missing) or '工具/方法/结果'}，写进项目经历而不是只放技能栏。"
        elif score <= 7:
            suggestion = f"{name} 基础可用，下一步补充量化结果或真实场景。"
        radar.append({"category": name, "score": score, "matched": matched, "missing": missing, "suggestion": suggestion})
    return jsonify({"success": True, "radar_data": radar, "ai_comment": "技能图谱已根据简历关键词生成，建议补足低分象限的项目证据。", "ai_used": False})


@app.route("/api/resume-generator", methods=["POST"])
def resume_generator():
    data = request.get_json() or {}
    name = data.get("name", "候选人")
    target = data.get("job_target", "目标岗位")
    skills = data.get("skills", "Python, Flask, Selenium, JMeter")
    content = f"""{name}
求职意向：{target}

核心技能：{skills}

项目经历：AI 求职辅助 Web 系统
- 负责系统需求梳理、功能实现与测试验证，覆盖简历管理、JD 匹配、模拟面试、求职进度看板。
- 设计接口测试、功能测试和性能测试用例，输出测试报告与缺陷记录。
- 通过多模型 API 接入和本地兜底策略，提高 AI 功能可用性。

自我评价：学习能力强，能把课程实训、真实求职场景和工程实现结合起来，重视可用性与测试闭环。
"""
    return jsonify({"success": True, "resume_content": content})


@app.route("/api/ai/analyze-jd", methods=["POST"])
def analyze_jd():
    data = request.get_json() or {}
    jd = data.get("jd_content", "")
    keywords = extract_keywords(jd)
    focus = extract_jd_focus(jd)
    risk_flags = []
    if any(word in jd for word in ["抗压", "高强度", "能加班", "狼性"]):
        risk_flags.append("JD 中出现高强度/加班暗示，面试时建议确认工作节奏。")
    if not keywords:
        risk_flags.append("JD 信息较少，建议补充岗位职责和任职要求后再分析。")
    return jsonify({
        "success": True,
        "content": "## JD 解析\n"
        f"- 核心关键词：{', '.join(keywords) or '需补充 JD'}\n"
        f"- 硬技能：{', '.join(focus.get('硬技能', [])) or '未明显出现'}\n"
        f"- 测试能力：{', '.join(focus.get('测试能力', [])) or '未明显出现'}\n"
        f"- 风险提示：{'；'.join(risk_flags) if risk_flags else '暂未发现明显风险词'}\n"
        "- 面试准备：准备一个项目深挖案例、一个问题定位案例、一个协作沟通案例。\n"
        "- 简历策略：把 JD 高频词写入项目经历，而不是只堆在技能栏。",
        "keywords": keywords,
        "focus": focus,
        "risk_flags": risk_flags,
        "ai_used": False,
    })


@app.route("/api/ai/compare-jds", methods=["POST"])
def compare_jds():
    data = request.get_json() or {}
    jds = data.get("jds", [])
    summaries = []
    for index, jd in enumerate(jds, 1):
        summaries.append(f"JD{index}：关键词 {', '.join(extract_keywords(jd)[:6]) or '不明显'}")
    return jsonify({"success": True, "content": "## 多 JD 对比\n" + "\n".join(f"- {item}" for item in summaries) + "\n\n建议优先选择关键词与你项目经历重合度最高的岗位。", "ai_used": False})


@app.route("/api/ai/generate-test-report", methods=["POST"])
def generate_test_report():
    data = request.get_json() or {}
    info = data.get("project_info", "AI 求职辅助 Web 系统")
    return jsonify({
        "success": True,
        "content": f"## 测试总结报告\n项目：{info}\n\n### 测试范围\n简历管理、JD 匹配、模拟面试、AI 助手、求职看板。\n\n### 结论\n核心流程可测，建议继续补充接口自动化、浏览器兼容性和异常上传用例。",
        "ai_used": False,
    })


@app.route("/api/resume-templates")
def resume_templates():
    return jsonify({
        "success": True,
        "data": {
            "campus_test": {
                "name": "应届测试工程师模板",
                "description": "突出测试工具、课程项目、缺陷报告和自动化意识。",
                "sections": ["个人信息", "求职意向", "核心技能", "项目经历", "测试实践", "教育背景"],
                "tips": ["每个项目写清测试对象", "补充工具和指标", "保留实训报告作为证据"],
            },
            "ai_product": {
                "name": "AI 应用项目模板",
                "description": "适合把本系统包装成 AI Agent 项目经历。",
                "sections": ["项目背景", "技术架构", "智能体能力", "模型接入", "测试验证", "项目成果"],
                "tips": ["强调多模型路由", "强调兜底策略", "强调真实求职流程"],
            },
        },
    })


if __name__ == "__main__":
    init_db()
    print("JobHunter AI running at http://localhost:5000")
    print("Providers: GLM / DeepSeek / Kimi, with local fallback.")
    app.run(debug=True, host="0.0.0.0", port=5000)
